"""Expand user uploads (HEIC, PDF, Excel, etc.) into shapes vLLM can consume."""

from __future__ import annotations

import base64
import io
import logging
import re
from typing import Any

from PIL import Image

try:
    from pillow_heif import register_heif_opener

    register_heif_opener()
except ImportError:
    pass

logger = logging.getLogger(__name__)

# Per-document cap. With 64k context (~4 chars/token ≈ 256k chars),
# cap each doc at 80k so two documents + prompt + output all fit.
MAX_EXTRACTED_CHARS = 80_000


def _truncate(s: str) -> str:
    if len(s) <= MAX_EXTRACTED_CHARS:
        return s
    return s[: MAX_EXTRACTED_CHARS - 20] + "\n\n[… truncated …]"


def parse_data_url(url: str) -> tuple[bytes, str]:
    """Return raw bytes and MIME from a data URL."""
    if not isinstance(url, str) or not url.startswith("data:"):
        return b"", ""
    m = re.match(r"data:([^;,]+)?(;base64)?,(.*)", url, re.DOTALL)
    if not m:
        return b"", ""
    mime = (m.group(1) or "application/octet-stream").strip().lower()
    payload = m.group(3) or ""
    if m.group(2):
        raw = base64.standard_b64decode(payload)
    else:
        raw = payload.encode("utf-8", errors="replace")
    return raw, mime


def normalize_image_part(part: dict[str, Any]) -> dict[str, Any]:
    """Re-encode raster data URLs as PNG so vLLM gets a known-good format (HEIC, AVIF, etc.)."""
    url = (part.get("image_url") or {}).get("url") or ""
    if not url.startswith("data:"):
        return part
    raw, _mime = parse_data_url(url)
    if not raw:
        return part
    try:
        img = Image.open(io.BytesIO(raw))
        img.load()
        img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG", optimize=True)
        b64 = base64.standard_b64encode(buf.getvalue()).decode("ascii")
        return {
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{b64}"},
        }
    except Exception as e:
        logger.warning("Image normalize failed, passing through: %s", e)
        return part


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    chunks: list[str] = []
    for page in reader.pages:
        chunks.append((page.extract_text() or "").strip())
    out = "\n\n".join(chunks).strip()
    return out or "(No extractable text in this PDF — it may be scanned images only.)"


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    try:
        for name in wb.sheetnames:
            ws = wb[name]
            lines.append(f"## Sheet: {name}")
            for row in ws.iter_rows(values_only=True):
                line = "\t".join("" if v is None else str(v) for v in row)
                if line.strip():
                    lines.append(line)
    finally:
        wb.close()
    return "\n".join(lines).strip() or "(Empty spreadsheet.)"


def _extract_xls(data: bytes) -> str:
    import xlrd

    book = xlrd.open_workbook(file_contents=data)
    lines: list[str] = []
    for si in range(book.nsheets):
        sh = book.sheet_by_index(si)
        lines.append(f"## Sheet: {sh.name}")
        for ri in range(sh.nrows):
            row = sh.row(ri)
            cells = [str(c.value) if c.value != "" else "" for c in row]
            line = "\t".join(cells)
            if line.strip():
                lines.append(line)
    return "\n".join(lines).strip() or "(Empty spreadsheet.)"


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append("\t".join(cells))
    return "\n\n".join(paragraphs).strip() or "(No extractable text in this Word document.)"


def _extract_csv_or_text(data: bytes, mime: str) -> str:
    encs = ("utf-8", "utf-8-sig", "cp1252", "latin-1")
    for enc in encs:
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_document_text(data: bytes, filename: str, mime: str) -> str:
    name = (filename or "").lower()
    mt = (mime or "").lower()

    if "pdf" in mt or name.endswith(".pdf"):
        try:
            return _truncate(_extract_pdf(data))
        except Exception as e:
            logger.exception("PDF extraction failed")
            raise ValueError(f"Could not read this PDF: {e}") from e

    if name.endswith(".docx") or "wordprocessingml" in mt:
        try:
            return _truncate(_extract_docx(data))
        except Exception as e:
            logger.exception("DOCX extraction failed")
            raise ValueError(f"Could not read this Word file: {e}") from e

    if name.endswith(".doc") and not name.endswith(".docx"):
        if "msword" in mt or mt == "application/msword":
            raise ValueError(
                "Old .doc format is not supported. Please save as .docx and re-upload."
            )

    if name.endswith(".xlsx") or "spreadsheetml" in mt:
        try:
            return _truncate(_extract_xlsx(data))
        except Exception as e:
            logger.exception("XLSX extraction failed")
            raise ValueError(f"Could not read this Excel file (.xlsx): {e}") from e

    if name.endswith(".xls") and not name.endswith(".xlsx"):
        try:
            return _truncate(_extract_xls(data))
        except Exception as e:
            logger.exception("XLS extraction failed")
            raise ValueError(f"Could not read this Excel file (.xls): {e}") from e

    if name.endswith(".csv") or "csv" in mt or "text/plain" in mt or name.endswith(".txt") or name.endswith(".md"):
        return _truncate(_extract_csv_or_text(data, mime))

    raise ValueError(f"Unsupported document type: {filename} ({mime or 'unknown'})")


def expand_file_part(part: dict[str, Any]) -> list[dict[str, Any]]:
    """Turn a client file part into OpenAI-style text parts for the model."""
    f = part.get("file") or {}
    data_url = f.get("dataUrl") or f.get("data_url") or ""
    filename = f.get("filename") or f.get("name") or "attachment"
    mime = f.get("mime") or f.get("mime_type") or ""
    logger.info(
        "expand_file_part: filename=%s, mime=%s, dataUrl_len=%d",
        filename, mime, len(data_url),
    )
    raw, parsed_mime = parse_data_url(data_url)
    if not raw:
        logger.error("expand_file_part: empty bytes after parsing data URL for %s", filename)
        raise ValueError("Empty file upload.")
    mime = mime or parsed_mime
    logger.info("expand_file_part: decoded %d bytes, mime=%s", len(raw), mime)
    body = extract_document_text(raw, filename, mime)
    logger.info(
        "expand_file_part: extracted %d chars from %s", len(body), filename,
    )
    header = f"--- Extracted content from **{filename}** ---\n\n"
    return [{"type": "text", "text": _truncate(header + body)}]


def normalize_messages_for_llm(messages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Apply file expansion and image re-encoding for every message."""
    out: list[dict[str, Any]] = []
    for msg in messages:
        content = msg.get("content")
        if not isinstance(content, list):
            out.append(msg)
            continue
        part_types = [p.get("type") if isinstance(p, dict) else type(p).__name__ for p in content]
        logger.info("normalize: role=%s, %d parts, types=%s", msg.get("role"), len(content), part_types)
        new_parts: list[dict[str, Any]] = []
        for p in content:
            if not isinstance(p, dict):
                continue
            t = p.get("type")
            if t == "file":
                try:
                    new_parts.extend(expand_file_part(p))
                except ValueError as e:
                    logger.error("File expansion failed for part: %s", e)
                    new_parts.append(
                        {"type": "text", "text": f"[Attachment error: {e}]"}
                    )
            elif t == "image_url":
                new_parts.append(normalize_image_part(p))
            else:
                new_parts.append(p)
        out.append({**msg, "content": new_parts})
    return out
